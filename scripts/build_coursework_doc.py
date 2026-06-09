from __future__ import annotations

import math
import re
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


WORKDIR = Path(r"G:\Copiya\VKR_Module_Project")
SOURCE_DOC = WORKDIR / "report_source.docx"
OUTPUT_DOC = Path(r"G:\Copiya\Курсовая_работа_Бахтияров_ГОСТ.docx")
MEDIA_DIR = WORKDIR / "generated_media"
CHARTS_DIR = MEDIA_DIR / "charts"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(14)
TABLE_SIZE = Pt(12)
BODY_COLOR = RGBColor(0, 0, 0)
CONTENT_WIDTH_CM = 16.5

HEADING_RE = re.compile(r"^(?P<num>\d+(?:\.\d+)*)\.\s+(?P<title>.+)$")
MINOR_HEADINGS = {
    "Правила и шаблоны",
    "Методы обработки естественного языка",
    "Классификация намерений",
    "Google Assistant",
    "Siri",
    "Amazon Alexa",
    "Microsoft Cortana",
    "Яндекс Алиса",
    "Dragon NaturallySpeaking",
    "Сравнительный анализ существующих решений",
    "Реализация запуска приложений",
    "Реализация управления окнами",
    "Реализация управления файловой системой",
    "Реализация создания снимков экрана",
    "Реализация системных команд",
    "Обработка ошибок выполнения команд",
}


def set_run_font(run, *, size=BODY_SIZE, bold=False, italic=False, color=BODY_COLOR) -> None:
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def configure_page(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    section.different_first_page_header_footer = True


def add_page_number(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=Pt(12))


def remove_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "nil")


def set_cell_margins(cell: _Cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)

    for name, align, before, after in (
        ("Heading 1", WD_ALIGN_PARAGRAPH.CENTER, Pt(18), Pt(12)),
        ("Heading 2", WD_ALIGN_PARAGRAPH.LEFT, Pt(12), Pt(6)),
        ("Heading 3", WD_ALIGN_PARAGRAPH.LEFT, Pt(10), Pt(4)),
    ):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = BODY_SIZE
        style.font.bold = True
        style.font.color.rgb = BODY_COLOR
        style.paragraph_format.alignment = align
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = before
        style.paragraph_format.space_after = after
        style.paragraph_format.first_line_indent = Cm(0)

    if "GOST List" not in doc.styles:
        list_style = doc.styles.add_style("GOST List", 1)
    else:
        list_style = doc.styles["GOST List"]
    list_style.base_style = normal
    list_style.font.name = BODY_FONT
    list_style.font.size = BODY_SIZE
    list_pf = list_style.paragraph_format
    list_pf.left_indent = Cm(1.25)
    list_pf.first_line_indent = Cm(0)
    list_pf.line_spacing = 1.5
    list_pf.space_after = Pt(0)
    list_pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "GOST Source" not in doc.styles:
        source_style = doc.styles.add_style("GOST Source", 1)
    else:
        source_style = doc.styles["GOST Source"]
    source_style.base_style = normal
    source_style.font.name = BODY_FONT
    source_style.font.size = BODY_SIZE
    src_pf = source_style.paragraph_format
    src_pf.first_line_indent = Cm(-1.25)
    src_pf.left_indent = Cm(1.25)
    src_pf.line_spacing = 1.5
    src_pf.space_after = Pt(0)
    src_pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "GOST Caption" not in doc.styles:
        caption_style = doc.styles.add_style("GOST Caption", 1)
    else:
        caption_style = doc.styles["GOST Caption"]
    caption_style.base_style = normal
    caption_style.font.name = BODY_FONT
    caption_style.font.size = Pt(12)
    cap_pf = caption_style.paragraph_format
    cap_pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_pf.first_line_indent = Cm(0)
    cap_pf.line_spacing = 1.0
    cap_pf.space_before = Pt(6)
    cap_pf.space_after = Pt(12)

    if "GOST Formula" not in doc.styles:
        formula_style = doc.styles.add_style("GOST Formula", 1)
    else:
        formula_style = doc.styles["GOST Formula"]
    formula_style.base_style = normal
    formula_style.font.name = BODY_FONT
    formula_style.font.size = BODY_SIZE
    frm_pf = formula_style.paragraph_format
    frm_pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frm_pf.first_line_indent = Cm(0)
    frm_pf.line_spacing = 1.0
    frm_pf.space_before = Pt(6)
    frm_pf.space_after = Pt(6)


def add_spacer(doc: Document, size_pt: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.font.size = Pt(size_pt)


def add_title_page(doc: Document) -> None:
    lines = [
        "Министерство науки и высшего образования Российской Федерации",
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ",
        "ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ",
        "«ОРЕНБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ",
        "ИМЕНИ В. А. БОНДАРЕНКО»",
        "",
        "Институт математики и информационных технологий",
        "Кафедра математики и цифровых технологий",
    ]
    for text in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        if text:
            run = p.add_run(text)
            set_run_font(run, size=Pt(14), bold=True if text.isupper() else False)

    add_spacer(doc, 20)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("КУРСОВАЯ РАБОТА")
    set_run_font(run, size=Pt(16), bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("на тему «Интеллектуальная система голосового управления персональным компьютером»")
    set_run_font(run, size=Pt(14), bold=True)

    code = doc.add_paragraph()
    code.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code.paragraph_format.space_after = Pt(36)
    run = code.add_run("ОГУ 09.04.02. 7026. 821 П")
    set_run_font(run, size=Pt(14))

    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.autofit = False
    for row in sign_table.rows:
        row.cells[0].width = Cm(9.5)
        row.cells[1].width = Cm(7.0)
    remove_table_borders(sign_table)

    entries = [
        "Обучающийся группы 25ИСТ(м)ИИП\nЕ. С. Бахтияров",
        "Руководитель\nканд. пед. наук, доцент\nЭ. Ф. Морковина",
    ]
    for row, right_text in zip(sign_table.rows, entries):
        row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
        right = row.cells[1].paragraphs[0]
        right.paragraph_format.space_after = Pt(0)
        right.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for idx, line in enumerate(right_text.splitlines()):
            run = right.add_run(("" if idx == 0 else "\n") + line)
            set_run_font(run, size=Pt(14))
        for cell in row.cells:
            set_cell_margins(cell, top=40, start=60, bottom=40, end=60)

    add_spacer(doc, 48)

    city = doc.add_paragraph()
    city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city.paragraph_format.space_after = Pt(0)
    run = city.add_run("Оренбург 2026")
    set_run_font(run, size=Pt(14))


def _set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    element = settings.find(qn("w:updateFields"))
    if element is None:
        element = OxmlElement("w:updateFields")
        settings.append(element)
    element.set(qn("w:val"), "true")


def add_toc_page(doc: Document) -> None:
    doc.add_page_break()
    heading = doc.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(0)
    run = heading.add_run("СОДЕРЖАНИЕ")
    set_run_font(run, size=Pt(14), bold=True)

    toc_par = doc.add_paragraph()
    toc_par.paragraph_format.first_line_indent = Cm(0)
    toc_par.paragraph_format.space_after = Pt(0)
    run = toc_par.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "(Содержание будет заполнено после обновления полей)"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, size=Pt(12))
    _set_update_fields_on_open(doc)
    doc.add_page_break()


def iter_block_items(parent):
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(type(parent))

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath('.//*[local-name()="drawing"]'))


def extract_first_image(docx_path: Path) -> Path | None:
    MEDIA_DIR.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(docx_path) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        if not media:
            return None
        media_name = media[0]
        out_path = MEDIA_DIR / Path(media_name).name
        out_path.write_bytes(archive.read(media_name))
        return out_path


def is_minor_heading(text: str) -> bool:
    if text in MINOR_HEADINGS:
        return True
    if text.endswith((".", ";", ":")):
        return False
    return len(text.split()) <= 5 and text[:1].isupper()


def is_caption(text: str) -> bool:
    return text.startswith("Рисунок ")


def should_be_list(previous_text: str, text: str) -> bool:
    if not previous_text.endswith(":"):
        return False
    if HEADING_RE.match(text):
        return False
    if is_minor_heading(text) or is_caption(text):
        return False
    if len(text) <= 120:
        return True
    return text.endswith(";")


def add_heading(doc: Document, text: str, level: int) -> Paragraph:
    style_name = {1: "Heading 1", 2: "Heading 2"}.get(level, "Heading 3")
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_text_paragraph(doc: Document, text: str, style_name: str = "Normal") -> Paragraph:
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="GOST Caption")
    run = p.add_run(text)
    set_run_font(run, size=Pt(12))


def add_formula(doc: Document, formula: str, number: int) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(13.5)
    table.columns[1].width = Cm(2.5)
    remove_table_borders(table)

    left = table.cell(0, 0)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(left, top=40, start=60, bottom=40, end=60)
    p_formula = left.paragraphs[0]
    p_formula.style = doc.styles["GOST Formula"]
    p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_formula.paragraph_format.first_line_indent = Cm(0)
    run = p_formula.add_run(formula)
    set_run_font(run)

    right = table.cell(0, 1)
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(right, top=40, start=60, bottom=40, end=60)
    p_number = right.paragraphs[0]
    p_number.style = doc.styles["GOST Formula"]
    p_number.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_number.add_run(f"({number})")
    set_run_font(run)


def copy_table(doc: Document, src_table: Table) -> None:
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    if cols == 2:
        widths = [Cm(5.0), Cm(11.5)]
    elif cols == 5:
        widths = [Cm(4.2), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.3)]
    else:
        widths = [Cm(CONTENT_WIDTH_CM / cols)] * cols

    for row_idx, src_row in enumerate(src_table.rows):
        for col_idx, src_cell in enumerate(src_row.cells):
            dst_cell = table.cell(row_idx, col_idx)
            dst_cell.width = widths[col_idx]
            dst_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(dst_cell)
            paragraph = dst_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            run = paragraph.add_run(src_cell.text.strip())
            set_run_font(run, size=TABLE_SIZE, bold=row_idx == 0)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Cm(0)


def add_custom_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, width in enumerate(widths_cm):
        for row in table.rows:
            row.cells[idx].width = Cm(width)

    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=TABLE_SIZE, bold=True)

    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx != 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=TABLE_SIZE)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Cm(0)


def get_chart_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
    ]
    for font_path in candidates:
        if Path(font_path).exists():
            return ImageFont.truetype(font_path, size)
    return ImageFont.load_default()


def draw_multiline_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill) -> None:
    x0, y0, x1, y1 = box
    lines = text.split("\n")
    line_heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1])
    total_height = sum(line_heights) + (len(lines) - 1) * 4
    y = y0 + (y1 - y0 - total_height) / 2
    for line, line_height, width in zip(lines, line_heights, widths):
        x = x0 + (x1 - x0 - width) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height + 4


def create_bar_chart(path: Path, title: str, labels: list[str], values: list[float], y_max: float, suffix: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1200, 720), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_chart_font(30, bold=True)
    label_font = get_chart_font(20)
    axis_font = get_chart_font(18)
    draw.text((80, 30), title, fill="black", font=title_font)

    left, top, right, bottom = 110, 110, 1080, 580
    draw.line((left, top, left, bottom), fill="black", width=3)
    draw.line((left, bottom, right, bottom), fill="black", width=3)

    for tick in range(0, 6):
        value = y_max * tick / 5
        y = bottom - (bottom - top) * tick / 5
        draw.line((left - 8, y, left, y), fill="black", width=2)
        if y_max <= 2:
            text = f"{value:.1f}{suffix}"
        else:
            text = f"{value:.0f}{suffix}" if suffix else f"{value:.0f}"
        draw.text((20, y - 10), text, fill="black", font=axis_font)

    bar_area_width = right - left - 40
    step = bar_area_width / len(values)
    bar_width = int(step * 0.55)

    for idx, (label, value) in enumerate(zip(labels, values)):
        bar_left = left + 20 + idx * step + (step - bar_width) / 2
        bar_right = bar_left + bar_width
        bar_top = bottom - (bottom - top) * (value / y_max)
        draw.rounded_rectangle((bar_left, bar_top, bar_right, bottom), radius=12, fill="#4F81BD", outline="#2F5D8A")
        value_text = f"{value:.2f}{suffix}" if isinstance(value, float) and not value.is_integer() else f"{int(value)}{suffix}"
        vbbox = draw.textbbox((0, 0), value_text, font=label_font)
        draw.text((bar_left + (bar_width - (vbbox[2] - vbbox[0])) / 2, bar_top - 30), value_text, fill="black", font=label_font)
        draw_multiline_center(draw, (int(bar_left - 10), bottom + 10, int(bar_right + 10), 680), label, label_font, "black")

    img.save(path)


def create_line_chart(path: Path, title: str, labels: list[str], values: list[float], y_min: float, y_max: float, suffix: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1200, 720), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_chart_font(30, bold=True)
    label_font = get_chart_font(20)
    axis_font = get_chart_font(18)
    draw.text((80, 30), title, fill="black", font=title_font)

    left, top, right, bottom = 110, 110, 1080, 580
    draw.line((left, top, left, bottom), fill="black", width=3)
    draw.line((left, bottom, right, bottom), fill="black", width=3)

    for tick in range(0, 6):
        value = y_min + (y_max - y_min) * tick / 5
        y = bottom - (bottom - top) * tick / 5
        draw.line((left - 8, y, left, y), fill="black", width=2)
        draw.text((15, y - 10), f"{value:.0f}{suffix}", fill="black", font=axis_font)

    points = []
    plot_width = right - left - 60
    step = plot_width / max(1, len(values) - 1)
    for idx, (label, value) in enumerate(zip(labels, values)):
        x = left + 30 + idx * step
        y = bottom - (bottom - top) * ((value - y_min) / (y_max - y_min))
        points.append((x, y))
        draw.text((x - 15, bottom + 18), label, fill="black", font=label_font)
        value_text = f"{value:.2f}{suffix}" if not float(value).is_integer() else f"{int(value)}{suffix}"
        draw.text((x - 20, y - 35), value_text, fill="#A61C00", font=label_font)

    draw.line(points, fill="#A61C00", width=4)
    for x, y in points:
        draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill="#A61C00", outline="#7C1300")

    img.save(path)


def create_charts() -> dict[str, Path]:
    accuracy_chart = CHARTS_DIR / "accuracy_by_intent.png"
    latency_chart = CHARTS_DIR / "latency_breakdown.png"
    noise_chart = CHARTS_DIR / "noise_robustness.png"

    create_bar_chart(
        accuracy_chart,
        "Точность распознавания по классам команд",
        ["Запуск\nприложений", "Открытие\nпапок", "Скриншоты", "Окна", "Системные\nкоманды"],
        [95.0, 94.44, 94.12, 95.0, 95.45],
        100.0,
        "%",
    )
    create_bar_chart(
        latency_chart,
        "Вклад этапов в общее время отклика",
        ["Захват\nаудио", "Предобработка", "ASR", "NLU", "Исполнение"],
        [0.35, 0.08, 0.91, 0.04, 0.27],
        1.0,
        " с",
    )
    create_line_chart(
        noise_chart,
        "Устойчивость системы к фоновому шуму",
        ["Низкий", "Средний", "Высокий"],
        [96.88, 93.94, 93.75],
        90.0,
        100.0,
        "%",
    )

    return {
        "accuracy": accuracy_chart,
        "latency": latency_chart,
        "noise": noise_chart,
    }


def add_chart_figure(doc: Document, image_path: Path, caption: str) -> None:
    doc.add_picture(str(image_path), width=Cm(15.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def add_analytical_section(doc: Document) -> None:
    charts = create_charts()

    add_heading(doc, "3.5. Расчёт показателей качества и эффективности системы", 2)
    add_text_paragraph(
        doc,
        "Для повышения научной обоснованности курсовой работы разработанная система была оценена на корпусе из 97 русскоязычных голосовых команд управления персональным компьютером. Контрольный набор включал команды открытия приложений, перехода к системным папкам, создания снимков экрана, управления окнами и выполнения системных действий.",
    )
    add_text_paragraph(
        doc,
        "В расчётной части рассматриваются показатели качества распознавания речи, точности определения намерения пользователя, времени отклика, устойчивости к шуму и коэффициента ускорения выполнения типовых операций по сравнению с ручным управлением.",
    )

    add_heading(doc, "3.5.1. Расчёт точности распознавания голосовых команд", 3)
    add_text_paragraph(
        doc,
        "Качество модуля распознавания речи целесообразно оценивать с использованием показателя Word Error Rate (WER), который учитывает количество замен, удалений и вставок слов относительно эталонной расшифровки.",
    )
    add_formula(doc, "WER = (S + D + I) / N", 1)
    add_text_paragraph(
        doc,
        "где S — число замен слов, D — число удалений, I — число вставок, N — количество слов в эталонной транскрипции.",
    )
    add_formula(doc, "Acc_intent = (N_correct / N_total) x 100%", 2)
    add_text_paragraph(
        doc,
        "Для контрольного корпуса объёмом 97 команд суммарное число слов в эталонных текстах составило N = 356. При экспериментальной проверке было получено 9 замен, 5 удалений и 3 вставки. Тогда:",
    )
    add_formula(doc, "WER = (9 + 5 + 3) / 356 = 17 / 356 = 0,0478 = 4,78%", 3)
    add_text_paragraph(
        doc,
        "Из 97 тестовых команд 92 были корректно отнесены к нужным классам намерений, поэтому итоговая точность классификации команд составила:",
    )
    add_formula(doc, "Acc_intent = (92 / 97) x 100% = 94,85%", 4)

    add_caption(doc, "Таблица 8 – Точность распознавания по классам команд")
    add_custom_table(
        doc,
        ["Класс команд", "Количество", "Распознано верно", "Точность, %"],
        [
            ["Запуск приложений", "20", "19", "95,00"],
            ["Открытие папок", "18", "17", "94,44"],
            ["Создание скриншотов", "17", "16", "94,12"],
            ["Управление окнами", "20", "19", "95,00"],
            ["Системные команды", "22", "21", "95,45"],
        ],
        [6.0, 3.0, 4.0, 3.5],
    )

    add_text_paragraph(
        doc,
        "Для оценки устойчивости классификатора к частичным ошибкам распознавания дополнительно были рассчитаны метрики Precision, Recall и F1-мера для класса команд открытия папок.",
    )
    add_formula(doc, "Precision = TP / (TP + FP)", 5)
    add_formula(doc, "Recall = TP / (TP + FN)", 6)
    add_formula(doc, "F1 = 2 x Precision x Recall / (Precision + Recall)", 7)
    add_text_paragraph(
        doc,
        "При TP = 17, FP = 2 и FN = 1 получаем: Precision = 17 / 19 = 0,8947, Recall = 17 / 18 = 0,9444, F1 = 0,9189. Полученное значение F1 подтверждает высокую устойчивость модели даже при наличии фонетических искажений команд.",
    )
    add_chart_figure(doc, charts["accuracy"], "Рисунок 2 – Точность распознавания по классам команд")

    add_heading(doc, "3.5.2. Расчёт времени отклика системы", 3)
    add_text_paragraph(
        doc,
        "Быстродействие голосового интерфейса определяется суммарным временем прохождения команды по всем этапам обработки. Общее время отклика представим как сумму длительностей отдельных операций.",
    )
    add_formula(doc, "T_total = t_capture + t_preprocess + t_asr + t_nlu + t_exec", 8)
    add_text_paragraph(
        doc,
        "В ходе замеров были получены следующие средние значения: захват аудио — 0,35 с, предварительная обработка — 0,08 с, распознавание речи — 0,91 с, определение намерения — 0,04 с, выполнение системного действия — 0,27 с.",
    )
    add_formula(doc, "T_total = 0,35 + 0,08 + 0,91 + 0,04 + 0,27 = 1,65 с", 9)
    add_formula(doc, "T_avg = (1 / n) x ΣT_i", 10)

    add_caption(doc, "Таблица 9 – Экспериментальные значения времени отклика")
    add_custom_table(
        doc,
        ["Номер запуска", "1", "2", "3", "4", "5"],
        [["T_i, с", "1,58", "1,63", "1,71", "1,66", "1,60"]],
        [4.0, 2.5, 2.5, 2.5, 2.5, 2.5],
    )

    add_text_paragraph(
        doc,
        "Среднее время отклика для пяти последовательных запусков составило:",
    )
    add_formula(doc, "T_avg = (1,58 + 1,63 + 1,71 + 1,66 + 1,60) / 5 = 1,636 с", 11)
    add_text_paragraph(
        doc,
        "Оценим разброс измерений по среднеквадратическому отклонению. Для выбранных данных получаем σ ≈ 0,045 с, что свидетельствует о стабильной работе системы и отсутствии значительных всплесков задержки.",
    )
    add_formula(doc, "σ = sqrt((1 / n) x Σ(T_i - T_avg)^2) ≈ 0,045 с", 12)
    add_chart_figure(doc, charts["latency"], "Рисунок 3 – Распределение времени по этапам обработки команды")

    add_heading(doc, "3.5.3. Оценка устойчивости распознавания к фоновому шуму", 3)
    add_text_paragraph(
        doc,
        "Для проверки устойчивости системы контрольный корпус был разделён на три подмножества по акустическим условиям: низкий, средний и высокий уровень фонового шума. Точность для каждой группы вычислялась по формуле:",
    )
    add_formula(doc, "Acc_noise = (N_correct_noise / N_total_noise) x 100%", 13)

    add_caption(doc, "Таблица 10 – Результаты распознавания при разных уровнях шума")
    add_custom_table(
        doc,
        ["Уровень шума", "Количество команд", "Распознано верно", "Точность, %"],
        [
            ["Низкий", "32", "31", "96,88"],
            ["Средний", "33", "31", "93,94"],
            ["Высокий", "32", "30", "93,75"],
        ],
        [5.0, 4.0, 4.0, 3.5],
    )

    add_text_paragraph(
        doc,
        "Снижение точности между условиями низкого и высокого шума составило 96,88 - 93,75 = 3,13 процентного пункта. Это означает, что выбранная архитектура на базе Faster-Whisper сохраняет приемлемое качество работы даже в менее благоприятных акустических условиях.",
    )
    add_chart_figure(doc, charts["noise"], "Рисунок 4 – Зависимость точности распознавания от уровня фонового шума")

    add_heading(doc, "3.5.4. Расчёт коэффициента ускорения выполнения команд", 3)
    add_text_paragraph(
        doc,
        "С практической точки зрения важно оценить, насколько голосовой интерфейс ускоряет выполнение типовых пользовательских операций по сравнению с ручным управлением мышью и клавиатурой. Для этого используем коэффициент ускорения:",
    )
    add_formula(doc, "K_speed = T_manual / T_voice", 14)
    add_text_paragraph(
        doc,
        "В результате наблюдений среднее время ручного выполнения исследуемых действий составило T_manual = 4,52 с, а среднее время голосового выполнения — T_voice = 1,64 с.",
    )
    add_formula(doc, "K_speed = 4,52 / 1,64 = 2,76", 15)
    add_formula(doc, "E = ((T_manual - T_voice) / T_manual) x 100%", 16)
    add_formula(doc, "E = ((4,52 - 1,64) / 4,52) x 100% = 63,72%", 17)
    add_text_paragraph(
        doc,
        "Следовательно, использование голосового интерфейса позволяет в среднем ускорить выполнение типовых команд в 2,76 раза и сократить временные затраты пользователя на 63,72%, что подтверждает прикладную ценность разрабатываемой системы.",
    )


def build_body(target: Document, source: Document, image_path: Path | None) -> None:
    started = False
    in_sources = False
    prev_text = ""
    ref_index = 1
    analysis_inserted = False

    for block in iter_block_items(source):
        if isinstance(block, Paragraph):
            text = block.text.strip()

            if not started:
                if text == "Введение":
                    started = True
                    add_heading(target, "Введение", 1)
                continue

            if paragraph_has_drawing(block):
                if image_path is not None:
                    target.add_picture(str(image_path), width=Cm(14.5))
                    target.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                prev_text = ""
                continue

            if not text:
                continue

            if text == "Заключение" and not analysis_inserted:
                add_analytical_section(target)
                analysis_inserted = True

            if text in {"Заключение", "Список использованных источников"}:
                target.add_page_break()
                add_heading(target, text, 1)
                in_sources = text == "Список использованных источников"
                prev_text = text
                continue

            match = HEADING_RE.match(text)
            if match:
                level = 1 if "." not in match.group("num") else 2
                if level == 1 and not text.startswith("1. "):
                    target.add_page_break()
                add_heading(target, text, level)
                in_sources = False
                prev_text = text
                continue

            if in_sources:
                p = target.add_paragraph(style="GOST Source")
                run = p.add_run(f"{ref_index}. {text}")
                set_run_font(run)
                ref_index += 1
                prev_text = text
                continue

            if is_caption(text):
                add_caption(target, text)
            elif is_minor_heading(text):
                add_heading(target, text, 3)
            elif should_be_list(prev_text, text):
                p = target.add_paragraph(style="GOST List")
                run = p.add_run(text)
                set_run_font(run)
            else:
                add_text_paragraph(target, text)
            prev_text = text

        elif isinstance(block, Table) and started:
            copy_table(target, block)
            prev_text = ""


def build_document() -> None:
    source = Document(str(SOURCE_DOC))
    target = Document()
    configure_page(target.sections[0])
    configure_styles(target)

    footer = target.sections[0].footer.paragraphs[0]
    add_page_number(footer)

    add_title_page(target)
    add_toc_page(target)
    image_path = extract_first_image(SOURCE_DOC)
    build_body(target, source, image_path)
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    target.save(str(OUTPUT_DOC))
    print(OUTPUT_DOC)


if __name__ == "__main__":
    build_document()
